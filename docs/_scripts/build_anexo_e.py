# -*- coding: utf-8 -*-
"""Genera el Anexo E — Guion del Video Demo Reel (.docx) con formato CIBERTEC."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

GREEN = RGBColor(0x2E, 0x7D, 0x32)

doc = Document()

# Formato base: Arial 11, A4, margenes manual (sup/inf 3cm, izq/der 2.5cm)
normal = doc.styles['Normal']
normal.font.name = 'Arial'
normal.font.size = Pt(11)
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
sec.top_margin = Cm(3)
sec.bottom_margin = Cm(3)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)


def h(text, size=14):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = GREEN
    r.font.name = 'Arial'
    return p


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    return p


# Titulo
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Anexo E — Guion del Video Demo Reel')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = GREEN
r.font.name = 'Arial'
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run('ReciclaPe — Desarrollo de Aplicaciones Web II (4697)')
rs.italic = True
rs.font.name = 'Arial'
rs.font.size = Pt(11)

para('')
para('Duración objetivo: 3 a 5 minutos. El video NO reemplaza la sustentación (manual §5.7). '
     'Recomendación de grabación: 1080p, screen-recording con OBS Studio, voz en off clara, '
     'subtítulos opcionales. Mostrar siempre la URL/acción en pantalla.')

# Ficha tecnica
h('1. Ficha técnica del video', 13)
ficha = [
    ('Formato', 'MP4 (H.264), 1920×1080, 30 fps'),
    ('Duración', '3:00 – 5:00 min'),
    ('Audio', 'Narración en off + música de fondo suave (libre de derechos)'),
    ('Herramientas sugeridas', 'OBS Studio (captura), Canva / CapCut (edición), intro con logo'),
    ('Entregable', 'ReciclaPe-DemoReel.mp4'),
]
tb = doc.add_table(rows=0, cols=2)
tb.style = 'Light Grid Accent 1'
tb.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in ficha:
    cells = tb.add_row().cells
    cells[0].text = k
    cells[1].text = v
    cells[0].paragraphs[0].runs[0].bold = True

para('')

# Guion por escenas
h('2. Guion por escenas (storyboard)', 13)

escenas = [
    ('00:00 – 00:20', 'Intro / Portada',
     'Logo ReciclaPe + título y nombres del equipo.',
     '"ReciclaPe es una plataforma de gestión de reciclaje vecinal construida con una arquitectura '
     'de microservicios. Conecta vecinos con recicladores formales y mide el impacto ambiental generado."'),
    ('00:20 – 00:45', 'El problema y la solución',
     'Datos clave en pantalla: 9 343 t/día de residuos en Lima, solo 3.9% reciclado. '
     'Esquema simple vecino → reciclador.',
     '"En Lima se generan más de 9 mil toneladas de residuos al día y se recicla menos del 4%. '
     'ReciclaPe reduce la fricción entre quien segrega y quien recicla."'),
    ('00:45 – 01:15', 'Arquitectura',
     'Mostrar el diagrama de arquitectura (Anexo A): gateway, eureka, auth, recojo, reporte, '
     'RabbitMQ, Prometheus/Grafana.',
     '"El backend son microservicios Spring Boot tras un API Gateway: autenticación con JWT y BCrypt, '
     'gestión de recojos, y reportes que se comunican por Feign y RabbitMQ."'),
    ('01:15 – 01:45', 'Login (rol Vecino)',
     'Pantalla de login. Ingresar credenciales. Mostrar brevemente el token JWT en DevTools / Network.',
     '"El login valida el password cifrado con BCrypt y emite un JWT firmado que el frontend usa '
     'en cada petición."'),
    ('01:45 – 02:30', 'Flujo del Vecino',
     'Crear un punto de acopio (CRUD), solicitar un recojo, ver el historial.',
     '"El vecino registra sus puntos de acopio y solicita un recojo de sus residuos segregados."'),
    ('02:30 – 03:15', 'Flujo del Reciclador',
     'Login como reciclador. Ver recojos disponibles, aceptar uno, completarlo registrando '
     'kilogramos por tipo de residuo.',
     '"El reciclador ve los recojos disponibles, acepta uno y al completarlo registra los kilogramos '
     'recuperados por tipo de residuo. Esto dispara un evento en RabbitMQ."'),
    ('03:15 – 03:50', 'Dashboard de impacto (rol Admin)',
     'Login como administrador. Mostrar dashboard distrital: kg recuperados, CO₂ evitado, '
     'número de recojos. Mostrar gráfico.',
     '"El administrador municipal visualiza el impacto agregado: kilogramos recuperados, CO₂ evitado '
     'y recojos completados, calculados por el servicio de reportes."'),
    ('03:50 – 04:25', 'Observabilidad y despliegue',
     'Mostrar `docker compose ps` con los contenedores arriba, Eureka con servicios registrados, '
     'un dashboard de Grafana con métricas de latencia.',
     '"Todo el stack se levanta con un solo docker compose up. Spring Boot Actuator expone métricas '
     'a Prometheus y las visualizamos en Grafana."'),
    ('04:25 – 05:00', 'Cierre',
     'Resumen de cumplimiento (rúbrica + sílabo) y logo final con datos del equipo.',
     '"ReciclaPe demuestra microservicios, seguridad, resiliencia, contenerización y observabilidad. '
     'Gracias por ver nuestro proyecto."'),
]

g = doc.add_table(rows=1, cols=4)
g.style = 'Light Grid Accent 1'
hdr = g.rows[0].cells
for i, txt in enumerate(['Tiempo', 'Escena', 'En pantalla', 'Narración (voz en off)']):
    hdr[i].text = txt
    hdr[i].paragraphs[0].runs[0].bold = True
for tiempo, escena, pantalla, narr in escenas:
    c = g.add_row().cells
    c[0].text = tiempo
    c[1].text = escena
    c[2].text = pantalla
    c[3].text = narr
    c[1].paragraphs[0].runs[0].bold = True

para('')

h('3. Checklist antes de grabar', 13)
for item in [
    'Base de datos cargada con el script del Anexo B (datos semilla).',
    'Todos los contenedores levantados (docker compose up) y servicios registrados en Eureka.',
    'Usuarios de prueba listos: vecino, reciclador y admin (password ReciclaPe2026).',
    'Cerrar pestañas y notificaciones; ocultar datos sensibles.',
    'Ensayar la narración para no exceder 5 minutos.',
    'Exportar en 1080p y revisar audio antes de entregar.',
]:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Arial'; run.font.size = Pt(11)

out = os.path.join(os.path.dirname(__file__), '..', 'Anexo-E-Guion-Video.docx')
out = os.path.abspath(out)
doc.save(out)
print('Guardado:', out)
